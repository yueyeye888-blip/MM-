from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "做市表格实时监控系统_U盘迁移与新电脑启动操作说明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent=120) -> None:
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=INK, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_numbering(doc: Document, kind: str, num_id: int, abstract_id: int) -> None:
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "number" else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)


def add_number_instance(doc: Document, num_id: int, abstract_id: int) -> None:
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        set_font(p.add_run(text))
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    set_num(p, 42)
    set_font(p.add_run(text))
    return p


_ACTIVE_NUMBER_ID = 41
_NEXT_NUMBER_ID = 100
_NEED_NUMBER_RESTART = True


def numbered(doc: Document, text: str):
    global _ACTIVE_NUMBER_ID, _NEXT_NUMBER_ID, _NEED_NUMBER_RESTART
    if _NEED_NUMBER_RESTART:
        add_number_instance(doc, _NEXT_NUMBER_ID, 41)
        _ACTIVE_NUMBER_ID = _NEXT_NUMBER_ID
        _NEXT_NUMBER_ID += 1
        _NEED_NUMBER_RESTART = False
    p = doc.add_paragraph(style="Normal")
    set_num(p, _ACTIVE_NUMBER_ID)
    set_font(p.add_run(text))
    return p


def code(doc: Document, lines: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "202733")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for idx, line in enumerate(lines.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=9.5, color=WHITE, font="Menlo")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def callout(doc: Document, label: str, text: str, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "：")
    set_font(r, bold=True, color=DARK_BLUE if fill != PALE_RED else "9B1C1C")
    set_font(p.add_run(text))
    return table


def info_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), bold=True, color=DARK_BLUE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_idx % 2:
                set_cell_shading(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(value), size=10.5)
    set_table_geometry(table, widths, indent=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def heading(doc: Document, text: str, level=1):
    global _NEED_NUMBER_RESTART
    if level <= 2:
        _NEED_NUMBER_RESTART = True
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


_PAGE_BREAK_COUNT = 0


def page_break(doc: Document):
    """Keep only the cover break; later sections flow naturally."""
    global _PAGE_BREAK_COUNT
    _PAGE_BREAK_COUNT += 1
    if _PAGE_BREAK_COUNT == 1:
        doc.add_page_break()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    set_font(paragraph.add_run(" 页"), size=9, color=MUTED)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_numbering(doc, "number", 41, 41)
    add_numbering(doc, "bullet", 42, 42)

    header = section.header.paragraphs[0]
    header.text = "做市表格实时监控系统 · 迁移操作手册"
    set_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover: compact editorial manual cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(74)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("操作手册"), size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("做市表格实时监控系统"), size=28, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    set_font(p.add_run("U盘迁移与新电脑启动操作说明"), size=17, color=BLUE)
    callout(doc, "适用范围", "旧 Mac → 新 Mac；完整复制程序、业务表格和历史数据后，在新电脑完成一次初始化。", LIGHT_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("系统版本：V3.2\n文档版本：1.0\n更新日期：2026年8月17日"), size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("本系统为本地只读旁路监控，不写入、不保存、不关闭业务工作簿。"), size=10, color=MUTED)

    page_break(doc)
    heading(doc, "1. 先看结论", 1)
    body(doc, "可以通过 U 盘把系统迁移到另一台 Mac。建议将文件从 U 盘复制到新电脑本地磁盘后运行，不建议长期直接在 U 盘中运行。新电脑第一次使用仍需安装运行环境、修正旧电脑路径并重新录入 GPT API Key。")
    callout(doc, "推荐方式", "完整迁移：复制整个系统目录、全部业务 Excel 和数据目录，可以保留历史对话、检测任务、报告、模板及项目数据库。", LIGHT_BLUE)
    heading(doc, "1.1 迁移结果", 2)
    info_table(doc, ["内容", "是否可通过 U 盘迁移", "说明"], [
        ["系统程序", "可以", "复制整个 MM表格数据汇报 文件夹"],
        ["业务 Excel", "可以", "所有被监控的 .xlsx/.xlsm 都要复制"],
        ["历史对话与快照", "可以", "位于 data 目录的数据库中"],
        ["检测任务与报告", "可以", "主要位于 data/control.db"],
        ["自定义对话模板", "可以", "位于 data/chat_templates.json"],
        ["GPT API Key", "不可以直接迁移", "保存在旧 Mac 钥匙串，新电脑需重新填写"],
        ["WPS Bridge 注册", "不可以直接迁移", "新电脑需重新安装并注册加载项"],
    ], [2050, 2050, 5260])
    heading(doc, "1.2 必要条件", 2)
    bullet(doc, "本说明默认旧电脑和新电脑都是 Mac。")
    bullet(doc, "新电脑需要拥有管理员权限或至少能够安装 Python、Node.js 和 WPS。")
    bullet(doc, "业务表格结构应保持一致，工作表名称为“APR实时表”。")
    bullet(doc, "项目名称由“APR实时表!A4:C4”读取；合并单元格时读取锚点 A4。")
    bullet(doc, "GPT 功能需要网络；表格保存后的本地监控不依赖互联网。")
    callout(doc, "Windows 提示", "当前启动文件是 macOS 的 .command，不能在 Windows 直接运行。Windows 电脑需要单独适配启动脚本、路径和 WPS Bridge。", PALE_RED)

    page_break(doc)
    heading(doc, "2. 旧电脑：制作 U 盘迁移包", 1)
    heading(doc, "2.1 停止程序并固定数据", 2)
    numbered(doc, "停止所有正在进行的检测时段，确认报告已经生成。")
    numbered(doc, "关闭监控系统的终端窗口；如无法确认，可重启旧电脑后暂时不要再次启动监控。")
    numbered(doc, "在 WPS 中保存全部业务表格，然后完整退出 WPS。")
    numbered(doc, "确认 U 盘剩余空间充足，并等待正在同步或写入的数据完成。")
    callout(doc, "重要", "必须先停止程序和退出 WPS再复制，避免 SQLite 数据库或 Excel 正在写入时得到不完整副本。", PALE_YELLOW)
    heading(doc, "2.2 复制系统文件", 2)
    body(doc, "在 Finder 中找到当前系统目录：")
    code(doc, "/Users/xingxiu/Desktop/MM表格数据汇报")
    body(doc, "将整个“MM表格数据汇报”文件夹复制到 U 盘。不要只复制启动文件，也不要遗漏隐藏目录或子目录。")
    heading(doc, "2.3 复制业务表格", 2)
    body(doc, "将所有已注册和计划监控的 Excel 文件一并复制到 U 盘。建议建立清晰目录：")
    code(doc, "做市监控迁移包/\n├── MM表格数据汇报/\n└── 业务表格/\n    ├── APR_做市实时表_V3_美化版.xlsx\n    └── 其他项目.xlsx")
    heading(doc, "2.4 复制后核对", 2)
    info_table(doc, ["必须存在", "用途"], [
        ["启动APR监控.command、run.sh、requirements.txt", "启动系统和安装 Python 依赖"],
        ["app、static、wps-bridge", "后端、页面和 WPS 内存读取组件"],
        ["data", "历史快照、对话、项目登记、检测任务和模板"],
        ["backups", "系统生成的业务表备份"],
        ["config.json", "基础路径、端口和 GPT 模型配置"],
        ["全部业务 Excel", "监控数据源"],
    ], [3300, 6060])

    page_break(doc)
    heading(doc, "3. 新电脑：准备运行环境", 1)
    heading(doc, "3.1 安装软件", 2)
    numbered(doc, "安装 WPS Office for Mac，并至少启动一次。")
    numbered(doc, "从 Python 官方安装 Python 3.11 或 3.12，安装后重新打开终端。")
    numbered(doc, "安装 Node.js 的长期支持版（LTS）；只有使用 WPS 未保存内存监控时才是必需项，但建议安装。")
    numbered(doc, "确认网络可访问 Python 和 Node.js 的依赖下载服务。")
    heading(doc, "3.2 检查版本", 2)
    body(doc, "打开“终端”，逐行执行：")
    code(doc, "python3 --version\npython3 -m pip --version\nnode --version\nnpm --version")
    body(doc, "能显示版本号即可。推荐 Python 3.11/3.12 和 Node.js LTS。")
    heading(doc, "3.3 从 U 盘复制到本地", 2)
    numbered(doc, "在新电脑桌面创建“MM表格数据汇报”文件夹，或直接复制同名文件夹。")
    numbered(doc, "将“业务表格”文件夹复制到新电脑固定位置，例如桌面或文稿目录。")
    numbered(doc, "迁移完成后不要改变系统目录内部结构。")
    callout(doc, "不要直接在 U 盘运行", "U 盘断开、休眠、盘符变化或写入速度不足都可能造成监控中断和数据库异常。应先复制到新电脑本地磁盘。", PALE_YELLOW)
    heading(doc, "3.4 安装系统依赖", 2)
    body(doc, "假设系统位于新电脑桌面，在终端执行：")
    code(doc, 'cd "$HOME/Desktop/MM表格数据汇报"\npython3 -m pip install -r requirements.txt\nchmod +x run.sh "启动APR监控.command"\ncd wps-bridge\nnpm install')
    callout(doc, "安装失败", "若 pip 或 npm 显示权限、证书或网络错误，不要反复启动系统；先按第 9 章排查。", PALE_RED)

    page_break(doc)
    heading(doc, "4. 修正新电脑路径", 1)
    body(doc, "系统和项目登记中保存了旧电脑的绝对路径，例如 /Users/xingxiu/Desktop/...。新电脑用户名或存放位置不同，必须在首次启动前修正。")
    heading(doc, "4.1 修改 config.json", 2)
    numbered(doc, "在系统目录中找到 config.json，先复制一份 config.json.bak 作为备份。")
    numbered(doc, "用文本编辑器打开 config.json。")
    numbered(doc, "把 source_workbook_path 修改为新电脑上主业务表格的完整路径。")
    body(doc, "示例：")
    code(doc, '"source_workbook_path": "/Users/新用户名/Desktop/业务表格/APR_做市实时表_V3_美化版.xlsx"')
    heading(doc, "4.2 修改 data/projects.json", 2)
    body(doc, "如果要保留原项目和历史数据库，请编辑 data/projects.json，并将以下字段中的旧路径替换为新路径：")
    bullet(doc, "每个项目的 workbook_path：指向新电脑上对应 Excel。")
    bullet(doc, "每个项目的 database_path：指向新电脑系统目录内对应数据库。")
    body(doc, "常用替换方式：将 `/Users/旧用户名` 整体替换为 `/Users/新用户名`；如果目录位置也变化，应替换完整前缀。")
    callout(doc, "路径规则", "JSON 中必须使用英文双引号，不能删除逗号；路径和文件名必须与 Finder 中实际文件完全一致。", PALE_YELLOW)
    heading(doc, "4.3 不保留项目登记时", 2)
    body(doc, "如果只要全新开始、不需要历史项目关系，可以把 data/projects.json 移到一个安全备份目录后启动，再在页面“管理项目”中重新注册表格。不要直接删除数据库文件。")
    callout(doc, "历史恢复", "若需要完整保留历史对话和报告，优先修改 projects.json 原路径，不要随意新建同名项目，以免生成新的 project_id 和新数据库。", LIGHT_BLUE)

    page_break(doc)
    heading(doc, "5. 首次启动", 1)
    heading(doc, "5.1 启动系统", 2)
    numbered(doc, "双击系统根目录中的“启动APR监控.command”。")
    numbered(doc, "macOS 首次拦截时，右键该文件 → 打开 → 再确认打开。")
    numbered(doc, "等待浏览器自动打开本地页面。")
    body(doc, "如果浏览器没有自动打开，手动访问：")
    code(doc, "http://127.0.0.1:8765")
    body(doc, "系统会使用两个本机端口：Agent 为 8765，WPS Bridge 页面为 3889。")
    heading(doc, "5.2 注册和核对项目", 2)
    numbered(doc, "进入页面右上角“管理项目”。")
    numbered(doc, "确认每个项目显示的工作簿名、完整路径和工作表均正确。")
    numbered(doc, "不存在的旧路径应先修正或移除，再选择新电脑上的对应表格。")
    numbered(doc, "确认项目名称与工作表 APR实时表 的 A4:C4 内容一致。")
    numbered(doc, "如果同名表格有多个副本，必须根据完整路径判断当前真正监控的是哪一个。")
    heading(doc, "5.3 验证保存后实时监控", 2)
    numbered(doc, "用 WPS 打开已经注册的业务表格。")
    numbered(doc, "记住页面当前的“数据时间”和一个准备测试的数值。")
    numbered(doc, "在 WPS 中修改该数值并手动点击保存。")
    numbered(doc, "等待约 1～3 秒，确认页面数据和“数据时间”更新。")
    numbered(doc, "恢复测试值并再次保存，确认页面同步恢复。")
    callout(doc, "合格标准", "手动保存后页面能在数秒内更新，即 DISK_REALTIME 磁盘实时监控正常；这一功能不依赖 WPS Bridge。", LIGHT_BLUE)

    page_break(doc)
    heading(doc, "6. 恢复 GPT 数据对话", 1)
    body(doc, "GPT API Key 保存在旧电脑 macOS 钥匙串的 APRMonitorOpenAI 项中，不会写入系统目录，因此不能靠 U 盘迁移。")
    numbered(doc, "打开监控页面的“数据对话”。")
    numbered(doc, "进入“API 设置”。")
    numbered(doc, "填写 OpenAI API Key 并保存。")
    numbered(doc, "使用一个简短模板生成报告，确认能正常返回。")
    callout(doc, "账号说明", "ChatGPT Plus 与 OpenAI API 独立。新电脑使用 GPT 数据对话时，API 账户仍需开通 API 计费。", PALE_YELLOW)
    heading(doc, "6.1 历史记录核对", 2)
    bullet(doc, "切换每个项目，确认历史问题与回答仍然存在。")
    bullet(doc, "检查自定义模板是否保留；模板文件为 data/chat_templates.json。")
    bullet(doc, "检查最近检测任务、时段和合并报告是否可打开。")
    bullet(doc, "若历史为空，优先检查 data 目录是否完整、projects.json 的 project_id 和 database_path 是否仍对应原数据库。")
    heading(doc, "6.2 数据口径抽查", 2)
    bullet(doc, "所有名称含“均价”的结果应固定保留 5 位小数。")
    bullet(doc, "现货持仓总均价应直接读取 APR实时表!G7:H7 合并区域的锚点 G7。")
    bullet(doc, "阶段买入均价应来自检测阶段内的数量和资金变化计算，不能使用当前市价或现货持仓总均价替代。")

    page_break(doc)
    heading(doc, "7. 安装并验证 WPS Bridge（可选）", 1)
    body(doc, "只有需要捕获“尚未保存到磁盘”的 WPS 内存值时，才必须安装 Bridge。仅依赖手动保存后的监控，可以跳过本章。")
    heading(doc, "7.1 注册加载项", 2)
    body(doc, "确认 Agent 正在运行，然后打开新终端执行：")
    code(doc, 'cd "$HOME/Desktop/MM表格数据汇报/wps-bridge"\nnpm install\nnpx wpsjs debug -p 3889')
    numbered(doc, "命令执行完成后，完整退出 WPS。")
    numbered(doc, "重新打开 WPS 和要监控的业务表格。")
    numbered(doc, "再次双击“启动APR监控.command”。")
    heading(doc, "7.2 检查状态", 2)
    body(doc, "浏览器访问：")
    code(doc, "http://127.0.0.1:8765/api/v1/bridge/status")
    body(doc, "以下四项同时满足才算未保存值通道生效：")
    info_table(doc, ["字段", "正确值", "含义"], [
        ["connected", "true", "Bridge 页面与 Agent 已连接"],
        ["mode", "WPS_MEMORY", "当前正在使用 WPS 内存数据"],
        ["event_registered", "true", "SheetChange 事件注册成功"],
        ["last_error", "null", "没有持续错误"],
    ], [2700, 2200, 4460])
    heading(doc, "7.3 未保存值实测", 2)
    numbered(doc, "记录一个测试单元格原值。")
    numbered(doc, "在 WPS 中修改它，但不要保存。")
    numbered(doc, "确认监控页面已经显示新值。")
    numbered(doc, "撤销修改或恢复原值，避免污染正式数据。")
    callout(doc, "安全提示", "Bridge 只读取，不会写入、保存、关闭或锁定工作簿。不同 WPS Mac 版本仍应在正式使用前完成一次实机测试。", LIGHT_BLUE)

    page_break(doc)
    heading(doc, "8. 最终验收清单", 1)
    info_table(doc, ["检查项", "合格标准"], [
        ["系统启动", "双击启动文件后浏览器能打开 127.0.0.1:8765"],
        ["项目路径", "页面显示的文件名、完整路径、工作表均正确"],
        ["项目名称", "与 APR实时表!A4:C4 一致"],
        ["磁盘监控", "WPS 保存后 1～3 秒内数据时间和数值更新"],
        ["历史数据库", "历史对话、检测任务、报告可查看"],
        ["GPT", "API 设置成功，模板报告可生成"],
        ["WPS Bridge", "如启用，四项状态全部满足且未保存值实测通过"],
        ["备份", "backups 目录可见，手动备份功能可执行"],
    ], [3300, 6060])
    heading(doc, "8.1 建议完成一次完整业务测试", 2)
    numbered(doc, "选择一个项目并开始检测任务。")
    numbered(doc, "在 WPS 中完成一次小范围数据修改并保存。")
    numbered(doc, "停止当前时段并生成检测报告。")
    numbered(doc, "检查现货数量变化、可用资金变化、市值变化和阶段买入均价。")
    numbered(doc, "用数据对话模板生成同一时段报告，对照确定性检测报告。")

    heading(doc, "9. 常见故障排查", 1)
    info_table(doc, ["现象", "处理方法"], [
        ["双击后启动失败", "查看系统目录 logs/agent.log 和 logs/wps-bridge.log；确认依赖已安装"],
        ["提示 python3 不存在", "重新安装 Python 3.11/3.12，关闭并重新打开终端"],
        ["提示缺少 fastapi/uvicorn/openpyxl", "在系统根目录重新执行 python3 -m pip install -r requirements.txt"],
        ["Bridge 启动失败", "在 wps-bridge 目录重新执行 npm install，再检查 Node/npm 版本"],
        ["页面打开但数据不更新", "核对项目完整路径；确认修改的是被监控文件并已在 WPS 保存"],
        ["项目名称错误", "检查 APR实时表!A4:C4，保存表格后等待页面同步"],
        ["历史对话或报告为空", "检查 data 是否完整，以及 projects.json 中 project_id/database_path 是否对应原数据库"],
        ["GPT 无法回答", "重新填写 API Key，检查 API 余额和网络；Plus 订阅不能代替 API 计费"],
        ["端口被占用", "关闭重复启动的监控进程；确认 8765 和 3889 未被其他程序占用"],
    ], [3300, 6060])
    callout(doc, "回退办法", "出现无法判断的问题时，不要删除 data 或 backups。先停止程序，把当前系统目录整体复制一份，再进行路径修正或重新安装依赖。", PALE_RED)

    page_break(doc)
    heading(doc, "10. 日常使用与备份建议", 1)
    bullet(doc, "每次使用前确认顶部当前项目和完整工作簿路径，避免监控错表。")
    bullet(doc, "业务操作完成后手动保存 WPS，并确认页面数据时间更新。")
    bullet(doc, "结束检测任务后立即检查报告是否完整。")
    bullet(doc, "定期将整个系统目录和业务表格复制到另一块存储设备。")
    bullet(doc, "备份前必须停止程序并退出 WPS。")
    bullet(doc, "不要单独删除 .db、.db-wal、.db-shm 或 projects.json。")
    bullet(doc, "不要把 GPT API Key 写入 config.json、说明文档或共享 U 盘。")
    heading(doc, "10.1 推荐备份命名", 2)
    code(doc, "做市监控完整备份_YYYY-MM-DD/\n├── MM表格数据汇报/\n└── 业务表格/")
    heading(doc, "10.2 核心地址", 2)
    info_table(doc, ["用途", "地址"], [
        ["监控看板", "http://127.0.0.1:8765"],
        ["API 文档", "http://127.0.0.1:8765/docs"],
        ["Bridge 状态", "http://127.0.0.1:8765/api/v1/bridge/status"],
        ["Agent 日志", "MM表格数据汇报/logs/agent.log"],
        ["Bridge 日志", "MM表格数据汇报/logs/wps-bridge.log"],
    ], [3300, 6060])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("—— 操作说明结束 ——"), size=11, bold=True, color=BLUE)

    doc.core_properties.title = "做市表格实时监控系统 U盘迁移与新电脑启动操作说明"
    doc.core_properties.subject = "Mac 电脑之间通过 U 盘迁移系统的操作手册"
    doc.core_properties.author = "做市表格实时监控系统项目组"
    doc.core_properties.keywords = "做市监控,U盘迁移,Mac,WPS Bridge,操作手册"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
